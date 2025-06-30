import sys
import os
import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageEnhance, ImageOps
import io
import pytesseract
import re


def clean_toc_line(line):
    # Match: "Section Title.....12" or "Section Title .......... 12"
    match = re.match(r'^(.*?)(\.{5,}|\s+)+(\d+)$', line.strip())
    if match:
        title = match.group(1).strip()
        page = match.group(3).strip()
        return f"{title} .......... {page}"
    return line


def pdf_to_word(pdf_path, docx_path):
    document = Document()
    with pdfplumber.open(pdf_path) as pdf:
        print(f"Opened PDF: {pdf_path}, {len(pdf.pages)} pages")
        for page_num, page in enumerate(pdf.pages, 1):
            print(f"Processing page {page_num}/{len(pdf.pages)}")
            # Extract text with basic layout
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                print(f"  Extracting text from page {page_num}")
                for line in text.split('\n'):
                    document.add_paragraph(line)
            else:
                print(f"  No text found on page {page_num}, running OCR...")
                # Render page as image for OCR
                try:
                    page_image = page.to_image(resolution=300).original
                    # Preprocess: convert to grayscale and increase contrast
                    gray_image = ImageOps.grayscale(page_image)
                    enhancer = ImageEnhance.Contrast(gray_image)
                    enhanced_image = enhancer.enhance(2.0)  # Increase contrast
                    # Set Tesseract config for English and improved layout
                    custom_config = r'--oem 3 --psm 6'
                    ocr_text = pytesseract.image_to_string(
                        enhanced_image, lang='eng', config=custom_config)
                    if ocr_text.strip():
                        print(f"    OCR extracted text from page {page_num}")
                        for line in ocr_text.split('\n'):
                            # If the line contains a long run of dots, clean it
                            if re.search(r'\.{5,}', line):
                                cleaned = clean_toc_line(line)
                                document.add_paragraph(cleaned)
                            else:
                                document.add_paragraph(line)
                    else:
                        print(f"    OCR found no text on page {page_num}")
                except Exception as e:
                    print(f"    OCR failed on page {page_num}: {e}")
            # Extract images
            for img_index, img in enumerate(page.images):
                print(f"  Extracting image {img_index+1}/{len(page.images)} on page {page_num}")
                # Extract image bytes
                x0, top, x1, bottom = img["x0"], img["top"], img["x1"], img["bottom"]
                # Clamp bbox to page bbox
                page_x0, page_top, page_x1, page_bottom = page.bbox
                x0 = max(x0, page_x0)
                top = max(top, page_top)
                x1 = min(x1, page_x1)
                bottom = min(bottom, page_bottom)
                if x0 >= x1 or top >= bottom:
                    print(f"    Skipping invalid bbox for image {img_index+1}")
                    continue  # Skip invalid bbox
                try:
                    cropped = page.within_bbox((x0, top, x1, bottom)).to_image(resolution=300)
                    img_bytes = cropped.original.convert("RGB")
                    img_stream = io.BytesIO()
                    img_bytes.save(img_stream, format="PNG")
                    img_stream.seek(0)
                    document.add_picture(img_stream, width=Inches(4))
                    print(f"    Image {img_index+1} added to document")
                except Exception as e:
                    print(f"Warning: Could not extract image on page {page_num}, image {img_index}: {e}")
            if page_num != len(pdf.pages):
                document.add_page_break()
    print(f"Saving DOCX to {docx_path}")
    document.save(docx_path)
    print(f"Saved DOCX to {docx_path}")


def main():
    if len(sys.argv) != 3:
        print(f"Usage: python {os.path.basename(__file__)} input.pdf output.docx")
        sys.exit(1)
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    if not os.path.isfile(pdf_path):
        print(f"Error: File '{pdf_path}' does not exist.")
        sys.exit(1)
    pdf_to_word(pdf_path, docx_path)
    print(f"Converted '{pdf_path}' to '{docx_path}' successfully.")


if __name__ == "__main__":
    main() 